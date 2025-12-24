import streamlit as st
import re
import pdfplumber
from docx import Document
from PIL import Image
import pytesseract
from keybert import KeyBERT
from groq import Groq
from io import BytesIO
from docx import Document as DocxDocument
from fpdf import FPDF

# =========================
# PAGE CONFIG
# =========================
st.set_page_config(
    page_title="Syllabus → Industry Project Generator",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)

# =========================
# CONFIG
# =========================
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
client = Groq(api_key="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")  # <-- Replace with your Groq API Key
kw_model = KeyBERT()

# =========================
# SIDEBAR (INPUT & OPTIONS)
# =========================
st.sidebar.title("📂 Input Configuration")

# Input type
input_type = st.sidebar.selectbox(
    "Syllabus Input Type",
    ["Text", "PDF", "DOCX", "Image"]
)

# Project Generation Mode
st.sidebar.markdown("### ⚙️ Project Mode")
project_mode = st.sidebar.radio(
    "Choose Project Generation Mode",
    ["Flagship Project (Single)", "Multiple Projects (5–6 Ideas)"]
)

# Project Difficulty
st.sidebar.markdown("### 🎯 Project Difficulty")
difficulty = st.sidebar.selectbox(
    "Select Difficulty Level",
    ["Beginner", "Intermediate", "Advanced"]
)

# Role selection (expanded)
st.sidebar.markdown("### 🧑‍💻 Role-Based Output")
roles = [
    "Data Scientist", "AI Engineer", "Machine Learning Engineer", "Software Engineer",
    "Data Engineer", "Research Scientist", "Business Analyst", "Computer Vision Engineer",
    "NLP Engineer", "Robotics Engineer", "IoT/Edge AI Engineer", "Cybersecurity Analyst"
]
selected_role = st.sidebar.selectbox("Select your role", roles)

# Dark mode toggle (optional)
dark_mode = st.sidebar.checkbox("Enable Dark Mode", value=False)

# =========================
# INPUT HANDLERS
# =========================
def extract_text_input(text):
    return text.strip()

def extract_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            if page.extract_text():
                text += page.extract_text() + "\n"
    return text

def extract_docx(file):
    doc = Document(file)
    return "\n".join(p.text for p in doc.paragraphs)

def extract_image(file):
    image = Image.open(file)
    return pytesseract.image_to_string(image)

# =========================
# TEXT PREPROCESSING
# =========================
def clean_text(text):
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"\b(unit|module)\s*\d+\b", "", text, flags=re.I)
    return text.strip()

# =========================
# CONCEPT CANONICALIZATION
# =========================
def canonicalize_concept(concept):
    concept = concept.lower()
    if "autoencoder" in concept:
        return "Autoencoders"
    if "network" in concept:
        return "Neural Networks"
    if "deep learning" in concept or "learning deep" in concept:
        return "Deep Learning"
    return concept.title()

# =========================
# CONCEPT EXTRACTION
# =========================
def extract_concepts(text, top_n=6):
    keywords = kw_model.extract_keywords(
        text,
        keyphrase_ngram_range=(1, 2),
        stop_words="english"
    )
    concepts = []
    seen = set()
    for kw, score in keywords:
        canonical = canonicalize_concept(kw)
        if canonical not in seen:
            seen.add(canonical)
            concepts.append(canonical)
        if len(concepts) == top_n:
            break
    return concepts

# =========================
# PROJECT GENERATION
# =========================
def generate_projects(concepts, mode, role, difficulty):
    if mode == "Flagship Project (Single)":
        concepts_text = ", ".join(concepts)
        prompt = f"""
        You are an expert AI mentor.
        Generate ONE flagship, end-to-end project using these concepts: {concepts_text}.
        Frame it for the role: {role}.
        Difficulty level: {difficulty}.
        Include:
        1. Project Title
        2. Industry Problem
        3. Proposed Solution (2–3 lines)
        4. Tools / Tech Stack
        5. Real dataset suggestions
        Make it industry-grade and resume-ready.
        """
    else:
        concepts_text = "\n".join(concepts)
        prompt = f"""
        You are an expert AI mentor.
        Generate MULTIPLE projects (one per concept) using these concepts:
        {concepts_text}
        Frame each project for the role: {role}.
        Difficulty level: {difficulty}.
        Include:
        1. Project Title
        2. Industry Problem
        3. Proposed Solution (2–3 lines)
        4. Tools / Tech Stack
        5. Real dataset suggestions
        Make all projects industry-grade and resume-ready.
        """
    response = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[
            {"role": "system", "content": "You are a senior AI architect and product mentor."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.55,
        max_tokens=1000
    )
    return response.choices[0].message.content

# =========================
# EXPORT FUNCTIONS
# =========================
def download_pdf(content):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for line in content.split("\n"):
        pdf.multi_cell(0, 8, line)
    # Return as BytesIO
    pdf_bytes = pdf.output(dest='S').encode('latin-1')
    return BytesIO(pdf_bytes)

def download_docx(content):
    doc = DocxDocument()
    for line in content.split("\n"):
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# =========================
# MAIN APP
# =========================
# Header
st.markdown("## 🎓 Syllabus → Industry Project Generator")
st.markdown(
    "Transform academic syllabi into **industry-grade, resume-ready project ideas**."
)

# Syllabus input
raw_text = ""
st.markdown("### 📥 Provide Syllabus")
if input_type == "Text":
    user_text = st.text_area("Paste syllabus text", height=200)
    if user_text:
        raw_text = extract_text_input(user_text)
else:
    file = st.file_uploader("Upload syllabus file")
    if file:
        if input_type == "PDF":
            raw_text = extract_pdf(file)
        elif input_type == "DOCX":
            raw_text = extract_docx(file)
        elif input_type == "Image":
            raw_text = extract_image(file)

# Process input
if raw_text:
    cleaned_text = clean_text(raw_text)
    with st.expander("🔍 View Cleaned Syllabus Text"):
        st.write(cleaned_text)

    if st.button("🚀 Generate Project Ideas", use_container_width=True):
        with st.spinner("Generating projects..."):
            concepts = extract_concepts(cleaned_text)
            projects = generate_projects(concepts, project_mode, selected_role, difficulty)

        # Display
        col1, col2 = st.columns([1, 3])
        with col1:
            st.markdown("### 🧠 Extracted Concepts")
            for c in concepts:
                st.markdown(f"- **{c}**")

        with col2:
            st.markdown("### 📌 Generated Projects")
            st.markdown(projects)

            # Export buttons
            pdf_data = download_pdf(projects)
            docx_data = download_docx(projects)
            st.download_button("📄 Download as PDF", data=pdf_data, file_name="projects.pdf")
            st.download_button("📝 Download as DOCX", data=docx_data, file_name="projects.docx")

# =========================
# FOOTER
# =========================
st.markdown("---")
st.caption(
    "Built using KeyBERT + Groq LLM • Industry-grade project generator • Resume-ready ideas"
)
